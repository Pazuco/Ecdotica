"""
Tests para el módulo de procesamiento de archivos (ProcesadorDeArchivos).
Cubre TXT, PDF y DOCX con degradación suave cuando las librerías son opcionales.
"""

import os
import sys
import tempfile
import unittest
from unittest import mock

from procesamiento import archivos

SAMPLE_TXT = os.path.join(os.path.dirname(__file__), '..', 'src', 'samples', 'sample.txt')


class TestExtraerTxt(unittest.TestCase):
    """Tests para extracción de archivos de texto plano."""

    def setUp(self):
        self.procesador = archivos.ProcesadorDeArchivos()

    def test_extrae_contenido_conocido(self):
        texto = self.procesador._extraer_txt(SAMPLE_TXT)
        self.assertIn('Amor sin libertad', texto)

    def test_extrae_contenido_no_vacio(self):
        texto = self.procesador._extraer_txt(SAMPLE_TXT)
        self.assertGreater(len(texto), 100)

    def test_codificacion_utf8_caracteres_especiales(self):
        contenido = 'Cañón, ñoño y corazón: acentos y eñes.'
        with tempfile.NamedTemporaryFile(
            mode='w', suffix='.txt', encoding='utf-8', delete=False
        ) as f:
            f.write(contenido)
            ruta = f.name
        try:
            texto = self.procesador._extraer_txt(ruta)
            self.assertEqual(texto, contenido)
        finally:
            os.unlink(ruta)

    def test_archivo_inexistente_lanza_excepcion(self):
        with self.assertRaises((FileNotFoundError, OSError)):
            self.procesador._extraer_txt('/ruta/que/no/existe.txt')


class TestDispatcherExtraerTexto(unittest.TestCase):
    """Tests para el método público extraer_texto() que despacha por extensión."""

    def setUp(self):
        self.procesador = archivos.ProcesadorDeArchivos()

    def test_dispatch_txt(self):
        texto = self.procesador.extraer_texto(SAMPLE_TXT)
        self.assertIn('Amor sin libertad', texto)

    def test_extension_no_soportada_lanza_value_error(self):
        with self.assertRaises(ValueError):
            self.procesador.extraer_texto('documento.odt')

    def test_extension_no_soportada_mensaje_descriptivo(self):
        with self.assertRaises(ValueError) as ctx:
            self.procesador.extraer_texto('documento.epub')
        self.assertIn('no soportado', str(ctx.exception))

    def test_dispatch_pdf_delega_a_metodo_privado(self):
        with mock.patch.object(
            self.procesador, '_extraer_pdf', return_value='texto pdf'
        ) as mock_pdf:
            resultado = self.procesador.extraer_texto('manuscrito.pdf')
        mock_pdf.assert_called_once_with('manuscrito.pdf')
        self.assertEqual(resultado, 'texto pdf')

    def test_dispatch_docx_delega_a_metodo_privado(self):
        with mock.patch.object(
            self.procesador, '_extraer_docx', return_value='texto docx'
        ) as mock_docx:
            resultado = self.procesador.extraer_texto('manuscrito.docx')
        mock_docx.assert_called_once_with('manuscrito.docx')
        self.assertEqual(resultado, 'texto docx')


class TestExtraerPdf(unittest.TestCase):
    """Tests para extracción de archivos PDF."""

    def setUp(self):
        self.procesador = archivos.ProcesadorDeArchivos()

    def test_degradacion_sin_pypdf2(self):
        """Sin PyPDF2, _extraer_pdf devuelve mensaje informativo en lugar de crash."""
        with mock.patch.dict(sys.modules, {'PyPDF2': None}):
            resultado = self.procesador._extraer_pdf('cualquier.pdf')
        self.assertEqual(resultado, 'PyPDF2 no instalado')

    def test_degradacion_devuelve_string(self):
        """El resultado es siempre una cadena, nunca None ni excepción."""
        with mock.patch.dict(sys.modules, {'PyPDF2': None}):
            resultado = self.procesador._extraer_pdf('cualquier.pdf')
        self.assertIsInstance(resultado, str)

    def test_con_pypdf2_extrae_sin_error(self):
        """Con PyPDF2 disponible, _extraer_pdf procesa el archivo sin lanzar excepción."""
        try:
            import PyPDF2
            from PyPDF2 import PdfWriter
        except ImportError:
            self.skipTest('PyPDF2 no instalado — test omitido')

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            writer.write(f)
            ruta = f.name

        try:
            resultado = self.procesador._extraer_pdf(ruta)
            self.assertIsInstance(resultado, str)
        finally:
            os.unlink(ruta)

    def test_con_pypdf2_multiples_paginas(self):
        """Con PyPDF2, _extraer_pdf concatena el texto de todas las páginas."""
        try:
            from PyPDF2 import PdfWriter
        except ImportError:
            self.skipTest('PyPDF2 no instalado — test omitido')

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.add_blank_page(width=200, height=200)

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            writer.write(f)
            ruta = f.name

        try:
            resultado = self.procesador._extraer_pdf(ruta)
            # No lanza excepción y devuelve string (páginas en blanco = cadena vacía)
            self.assertIsInstance(resultado, str)
        finally:
            os.unlink(ruta)


class TestExtraerDocx(unittest.TestCase):
    """Tests para extracción de archivos DOCX."""

    def setUp(self):
        self.procesador = archivos.ProcesadorDeArchivos()

    def test_degradacion_sin_python_docx(self):
        """Sin python-docx, _extraer_docx devuelve mensaje informativo en lugar de crash."""
        with mock.patch.dict(sys.modules, {'docx': None}):
            resultado = self.procesador._extraer_docx('cualquier.docx')
        self.assertEqual(resultado, 'python-docx no instalado')

    def test_degradacion_devuelve_string(self):
        """El resultado es siempre una cadena, nunca None ni excepción."""
        with mock.patch.dict(sys.modules, {'docx': None}):
            resultado = self.procesador._extraer_docx('cualquier.docx')
        self.assertIsInstance(resultado, str)

    def test_con_python_docx_extrae_parrafos(self):
        """Con python-docx, _extraer_docx extrae el texto de los párrafos."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest('python-docx no instalado — test omitido')

        doc = Document()
        doc.add_paragraph('Primer párrafo del manuscrito.')
        doc.add_paragraph('Segundo párrafo del manuscrito.')

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            ruta = f.name
        doc.save(ruta)

        try:
            resultado = self.procesador._extraer_docx(ruta)
            self.assertIn('Primer párrafo del manuscrito.', resultado)
            self.assertIn('Segundo párrafo del manuscrito.', resultado)
        finally:
            os.unlink(ruta)

    def test_con_python_docx_resultado_es_string(self):
        """El resultado de _extraer_docx es siempre una cadena."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest('python-docx no instalado — test omitido')

        doc = Document()
        doc.add_paragraph('Texto de prueba.')

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            ruta = f.name
        doc.save(ruta)

        try:
            resultado = self.procesador._extraer_docx(ruta)
            self.assertIsInstance(resultado, str)
        finally:
            os.unlink(ruta)

    def test_con_python_docx_parrafos_separados_por_salto(self):
        """Cada párrafo DOCX se separa con salto de línea en el texto extraído."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest('python-docx no instalado — test omitido')

        doc = Document()
        doc.add_paragraph('Línea uno.')
        doc.add_paragraph('Línea dos.')

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            ruta = f.name
        doc.save(ruta)

        try:
            resultado = self.procesador._extraer_docx(ruta)
            self.assertIn('\n', resultado)
        finally:
            os.unlink(ruta)


if __name__ == '__main__':
    unittest.main()
